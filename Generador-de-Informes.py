#!/usr/bin/env python3
# InformesAlMayor_fixed.py
# Versión corregida y mejorada del script original proporcionado por el usuario.
# Basado en: InformesAlMayor.py. (referencia incluida en la conversación). :contentReference[oaicite:1]{index=1}

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from openpyxl import load_workbook
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os
import json
import threading
from datetime import datetime
import re
import glob
import platform
import math

# ---------------------------
# Utilidades generales
# ---------------------------

def open_path(path):
    """Abrir carpeta o archivo de forma cross-platform."""
    try:
        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":
            # macOS
            os.system(f'open "{path}"')
        else:
            # linux
            os.system(f'xdg-open "{path}"')
    except Exception:
        pass

# ---------------------------
# Funciones auxiliares (Word/Excel/Imágenes)
# ---------------------------

def aplicar_formato_documento(doc, font_config):
    """Aplicar formato de fuente a todo el documento (parcial: runs)."""
    try:
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                try:
                    run.font.name = font_config.get("paragraph_font", "Calibri")
                    run.font.size = Pt(font_config.get("paragraph_size", 11))
                except Exception:
                    pass

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            try:
                                run.font.name = font_config.get("table_font", "Calibri")
                                run.font.size = Pt(font_config.get("table_size", 11))
                            except Exception:
                                pass
    except Exception:
        # No romper el flujo por errores de formato.
        pass

def reemplazar_texto_global(doc, viejo, nuevo):
    """Reemplazar texto en párrafos y tablas (busca en runs para mantener formato cuando es posible)."""
    if viejo == "":
        return
    # Reemplazos en párrafos
    for p in doc.paragraphs:
        if viejo in p.text:
            # reconstruir runs: operación simple y segura
            full = p.text.replace(viejo, nuevo)
            # limpiar runs
            for r in list(p.runs):
                r.text = ""
            p.add_run(full)

    # Reemplazos en tablas
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                if viejo in celda.text:
                    texto = celda.text.replace(viejo, nuevo)
                    # limpiar y escribir
                    for rp in list(celda.paragraphs):
                        rp.clear()  # available in python-docx 0.8.11+; if fails, fallback
                    # fallback seguro:
                    try:
                        celda.text = texto
                    except Exception:
                        # última opción: sustituir runs
                        for r in celda.paragraphs[0].runs:
                            r.text = ""
                        celda.paragraphs[0].add_run(texto)

def listar_imagenes_doc(doc):
    """
    Listar imágenes en el documento (párrafos y tablas).
    Devuelve lista de dicts: {'run': run_obj, 'paragraph': p_obj, 'idx_global': i}
    """
    imagenes = []
    idx = 0
    # revisar párrafos
    for p in doc.paragraphs:
        for run in p.runs:
            try:
                xml = run._element.xml
                if '<w:drawing' in xml or '<w:pict' in xml:
                    imagenes.append({"run": run, "paragraph": p, "idx_global": idx})
                    idx += 1
            except Exception:
                continue

    # revisar tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        try:
                            xml = run._element.xml
                            if '<w:drawing' in xml or '<w:pict' in xml:
                                imagenes.append({"run": run, "paragraph": p, "idx_global": idx})
                                idx += 1
                        except Exception:
                            continue

    return imagenes

def reemplazar_imagen(run, ruta_imagen, fixed_height_cm=None):
    """
    Reemplaza la imagen contenida en 'run' por la imagen en ruta.
    run: objeto Run (python-docx)
    """
    if not os.path.exists(ruta_imagen):
        raise FileNotFoundError(f"Imagen no encontrada: {ruta_imagen}")

    # Borrar contenido del run
    try:
        # clear() no está siempre documentado; en muchas versiones funciona
        run.clear()
    except Exception:
        # alternativa: poner texto vacío en cada run
        run.text = ""

    # Insertar nueva imagen
    try:
        if fixed_height_cm:
            run.add_picture(ruta_imagen, height=Cm(float(fixed_height_cm)))
        else:
            run.add_picture(ruta_imagen)
    except Exception as e:
        # fallback: intentar sin height
        try:
            run.add_picture(ruta_imagen)
        except Exception as e2:
            raise Exception(f"No se pudo insertar la imagen {ruta_imagen}: {str(e2)}")

def extraer_dato_excel_mejorado(excel_path, hoja, celda, tipo, decimales_config):
    """
    Extrae un valor o calcula promedio según 'tipo' desde excel_path.
    - celda soporta: "A1", "A1:A10", "C5,E7,F9", "C5:E10".
    - tipo: "valor" (devuelve primera celda válida), "promedio" (media numérica).
    """
    wb = load_workbook(excel_path, data_only=True)
    if hoja not in wb.sheetnames:
        raise ValueError(f"No se encontró la hoja '{hoja}' en {os.path.basename(excel_path)}")

    ws = wb[hoja]

    def cell_value_safe(coord):
        try:
            v = ws[coord].value
            return v
        except Exception:
            return None

    if tipo == "valor":
        # una celda simple
        if ":" not in celda and "," not in celda:
            v = cell_value_safe(celda)
            result = v if v is not None else ""
        else:
            # lista de celdas separadas por comas: devolver la primera no-vacía
            parts = [c.strip() for c in celda.split(",")]
            result = ""
            for p in parts:
                if ":" in p:
                    # rango -> tomar primer valor no-nulo del rango
                    try:
                        rng = ws[p]
                        for row in rng:
                            # row puede ser tuple
                            if isinstance(row, tuple):
                                for cell in row:
                                    if cell.value is not None:
                                        result = cell.value
                                        break
                                if result != "":
                                    break
                            else:
                                if row.value is not None:
                                    result = row.value
                                    break
                    except Exception:
                        continue
                else:
                    val = cell_value_safe(p)
                    if val is not None:
                        result = val
                        break
        # aplicar decimales si numérico
        if isinstance(result, (int, float)) and decimales_config.get("usar_decimales_fijos", False):
            dec = int(decimales_config.get("cantidad_decimales", 1))
            result = round(result, dec)
        return result

    elif tipo == "promedio":
        valores = []
        # si es rango con ':'
        if ":" in celda:
            try:
                rng = ws[celda]
                # rng puede ser generator de tuples
                for row in rng:
                    if isinstance(row, tuple):
                        for cell in row:
                            if isinstance(cell.value, (int, float)):
                                valores.append(cell.value)
                    else:
                        if isinstance(row.value, (int, float)):
                            valores.append(row.value)
            except Exception as e:
                raise ValueError(f"Error procesando rango {celda}: {str(e)}")
        elif "," in celda:
            parts = [c.strip() for c in celda.split(",")]
            for p in parts:
                try:
                    v = ws[p].value
                    if isinstance(v, (int, float)):
                        valores.append(v)
                except Exception:
                    continue
        else:
            # única celda
            v = ws[celda].value
            if isinstance(v, (int, float)):
                valores.append(v)

        if not valores:
            promedio = 0
        else:
            promedio = sum(valores) / len(valores)
            if decimales_config.get("usar_decimales_fijos", False):
                dec = int(decimales_config.get("cantidad_decimales", 1))
                promedio = round(promedio, dec)
        return promedio
    else:
        raise ValueError(f"Tipo inválido: {tipo}")

# ---------------------------
# Clase principal de la app
# ---------------------------

class CalicataApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Generador de Informes por Calicata - v3.0 Pro (corregido)")
        # intentamos maximizar de forma segura
        try:
            self.root.state('zoomed')
        except Exception:
            pass
        self.root.minsize(1200, 700)

        # Estilo
        style = ttk.Style()
        try:
            style.theme_use('clam')
        except Exception:
            pass

        # Configuración por defecto
        self.config = {
            "docx_path": "",
            "excel_folder_1": "",
            "excel_folder_2": "",
            "output_folder": "",
            "mappings": [],  # {"encabezado","hoja","celda","tipo"}
            "text_replacements": [],  # list of tuples
            "imagenes_folder": "",
            "image_replacements": [],
            "fixed_image_height": 5.0,
            "font_config": {
                "paragraph_font": "Calibri",
                "paragraph_size": 11,
                "table_font": "Calibri",
                "table_size": 11
            },
            "archivo_config": {
                "nombre_base": "EMS CUSCO C-",
                "usar_sufijo": True,
                "sufijo_personalizado": ""
            },
            "imagen_config": {
                "usar_mapeo_automatico": False,
                "imagen_mapeos": []
            },
            "informe_config": {
                "tipo_informe": "individual",
                "consolidado_nombre": "Informe_Consolidado"
            },
            "decimales_config": {
                "usar_decimales_fijos": True,
                "cantidad_decimales": 1
            }
        }

        # Estado
        self.processing = False
        self.stop_processing_flag = False
        self.editing_item = None
        self.last_config_file = None

        # Construir GUI
        self.build_gui()
        self.center_window()
        self.log("🎉 Aplicación iniciada (versión corregida).")

    def center_window(self):
        self.root.update_idletasks()
        w = self.root.winfo_width()
        h = self.root.winfo_height()
        sw = self.root.winfo_screenwidth()
        sh = self.root.winfo_screenheight()
        x = max((sw - w) // 2, 0)
        y = max((sh - h) // 2, 0)
        try:
            self.root.geometry(f"{w}x{h}+{x}+{y}")
        except Exception:
            pass

    def build_gui(self):
        # Main frame
        main = ttk.Frame(self.root, padding=8)
        main.grid(row=0, column=0, sticky="nsew")
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main.columnconfigure(0, weight=2)
        main.columnconfigure(1, weight=1)
        main.rowconfigure(1, weight=1)

        # Title
        title = ttk.Frame(main)
        title.grid(row=0, column=0, columnspan=2, sticky="ew")
        ttk.Label(title, text="🏗️ Generador de Informes por Calicata - Pro", font=("Arial", 16, "bold")).pack(side="left")
        self.status_label = ttk.Label(title, text="● Listo", foreground="#27ae60")
        self.status_label.pack(side="right")

        # Left config panel (with canvas + scrollbar)
        left_frame = ttk.LabelFrame(main, text="Configuración", padding=8)
        left_frame.grid(row=1, column=0, sticky="nsew", padx=(0,10))
        left_frame.columnconfigure(0, weight=1)
        left_frame.rowconfigure(0, weight=1)

        # Canvas y scrollbars en left_frame
        canvas = tk.Canvas(left_frame)
        v_scrollbar = ttk.Scrollbar(left_frame, orient="vertical", command=canvas.yview)
        h_scrollbar = ttk.Scrollbar(left_frame, orient="horizontal", command=canvas.xview)

        canvas.configure(yscrollcommand=v_scrollbar.set, xscrollcommand=h_scrollbar.set)

        # Grid layout
        canvas.grid(row=0, column=0, sticky="nsew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        h_scrollbar.grid(row=1, column=0, sticky="ew")

        # Necesario para que canvas pueda expandirse
        left_frame.rowconfigure(0, weight=1)
        left_frame.columnconfigure(0, weight=1)

        # Contenedor interno
        inner = ttk.Frame(canvas)
        canvas.create_window((0, 0), window=inner, anchor="nw")

        # Ajustar scrollregion automáticamente
        def _on_configure(event):
            canvas.configure(scrollregion=canvas.bbox("all"))

        inner.bind("<Configure>", _on_configure)


        row = 0
        # Files selection
        ttk.Label(inner, text="Documento Word base:").grid(row=row, column=0, sticky="w", pady=4)
        self.docx_entry = ttk.Entry(inner, width=70)
        self.docx_entry.grid(row=row, column=1, padx=6, sticky="w")
        ttk.Button(inner, text="📂", width=3, command=self.select_docx).grid(row=row, column=2)
        row += 1

        ttk.Label(inner, text="Carpeta Excel 1:").grid(row=row, column=0, sticky="w", pady=4)
        self.excel_folder_entry_1 = ttk.Entry(inner, width=70)
        self.excel_folder_entry_1.grid(row=row, column=1, padx=6, sticky="w")
        ttk.Button(inner, text="📂", width=3, command=lambda: self.select_excel_folder(1)).grid(row=row, column=2)
        row += 1

        ttk.Label(inner, text="Carpeta Excel 2:").grid(row=row, column=0, sticky="w", pady=4)
        self.excel_folder_entry_2 = ttk.Entry(inner, width=70)
        self.excel_folder_entry_2.grid(row=row, column=1, padx=6, sticky="w")
        ttk.Button(inner, text="📂", width=3, command=lambda: self.select_excel_folder(2)).grid(row=row, column=2)
        row += 1

        ttk.Label(inner, text="Carpeta de salida:").grid(row=row, column=0, sticky="w", pady=4)
        self.output_folder_entry = ttk.Entry(inner, width=70)
        self.output_folder_entry.grid(row=row, column=1, padx=6, sticky="w")
        ttk.Button(inner, text="📂", width=3, command=self.select_output_folder).grid(row=row, column=2)
        row += 1

        ttk.Separator(inner, orient="horizontal").grid(row=row, column=0, columnspan=3, sticky="ew", pady=8)
        row += 1

        # Archivo output naming
        archivo_frame = ttk.LabelFrame(inner, text="Nombres de archivo", padding=6)
        archivo_frame.grid(row=row, column=0, columnspan=3, sticky="ew", pady=6)
        archivo_frame.columnconfigure(1, weight=1)
        ttk.Label(archivo_frame, text="Nombre base:").grid(row=0, column=0, sticky="w")
        self.nombre_base_entry = ttk.Entry(archivo_frame, width=40)
        self.nombre_base_entry.grid(row=0, column=1, sticky="w")
        self.nombre_base_entry.insert(0, self.config["archivo_config"]["nombre_base"])
        self.usar_sufijo_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(archivo_frame, text="Usar sufijo automático (01,02...)", variable=self.usar_sufijo_var, command=self.toggle_sufijo).grid(row=1, column=0, columnspan=2, sticky="w", pady=4)
        ttk.Label(archivo_frame, text="Sufijo personalizado:").grid(row=2, column=0, sticky="w")
        self.sufijo_entry = ttk.Entry(archivo_frame, width=20, state="disabled")
        self.sufijo_entry.grid(row=2, column=1, sticky="w")

        row += 1
        ttk.Separator(inner, orient="horizontal").grid(row=row, column=0, columnspan=3, sticky="ew", pady=8)
        row += 1

        # Tipo de informe
        informe_frame = ttk.LabelFrame(inner, text="Tipo de informe", padding=6)
        informe_frame.grid(row=row, column=0, columnspan=3, sticky="ew", pady=6)
        self.tipo_informe_var = tk.StringVar(value="individual")
        ttk.Radiobutton(informe_frame, text="Individual (un docx por calicata)", variable=self.tipo_informe_var, value="individual").grid(row=0, column=0, sticky="w")
        ttk.Radiobutton(informe_frame, text="Consolidado (un docx con todos)", variable=self.tipo_informe_var, value="consolidado").grid(row=1, column=0, sticky="w")
        ttk.Label(informe_frame, text="Nombre consolidado:").grid(row=0, column=1, sticky="w")
        self.consolidado_nombre_entry = ttk.Entry(informe_frame, width=25)
        self.consolidado_nombre_entry.grid(row=0, column=2, sticky="w")
        self.consolidado_nombre_entry.insert(0, self.config["informe_config"]["consolidado_nombre"])

        row += 1
        ttk.Separator(inner, orient="horizontal").grid(row=row, column=0, columnspan=3, sticky="ew", pady=8)
        row += 1

        # Decimales config
        dec_frame = ttk.LabelFrame(inner, text="Formato numérico", padding=6)
        dec_frame.grid(row=row, column=0, columnspan=3, sticky="ew", pady=6)
        self.usar_decimales_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(dec_frame, text="Limitar decimales", variable=self.usar_decimales_var).grid(row=0, column=0, sticky="w")
        ttk.Label(dec_frame, text="Cantidad:").grid(row=0, column=1, sticky="w")
        self.decimales_spin = ttk.Spinbox(dec_frame, from_=0, to=6, width=5)
        self.decimales_spin.set("1")
        self.decimales_spin.grid(row=0, column=2, sticky="w")

        row += 1
        ttk.Separator(inner, orient="horizontal").grid(row=row, column=0, columnspan=3, sticky="ew", pady=8)
        row += 1

        # Formato de fuentes
        fmt_frame = ttk.LabelFrame(inner, text="Formato de texto", padding=6)
        fmt_frame.grid(row=row, column=0, columnspan=3, sticky="ew", pady=6)
        ttk.Label(fmt_frame, text="Párrafos - Fuente:").grid(row=0, column=0, sticky="w")
        self.paragraph_font = ttk.Combobox(fmt_frame, values=["Calibri", "Times New Roman", "Arial", "Verdana", "Tahoma"], width=18, state="readonly")
        self.paragraph_font.grid(row=0, column=1, sticky="w")
        self.paragraph_font.set("Calibri")
        ttk.Label(fmt_frame, text="Tamaño:").grid(row=0, column=2, sticky="w")
        self.paragraph_size = ttk.Spinbox(fmt_frame, from_=8, to=24, width=5)
        self.paragraph_size.set("11")
        self.paragraph_size.grid(row=0, column=3, sticky="w")

        ttk.Label(fmt_frame, text="Tablas - Fuente:").grid(row=1, column=0, sticky="w")
        self.table_font = ttk.Combobox(fmt_frame, values=["Calibri", "Times New Roman", "Arial", "Verdana", "Tahoma"], width=18, state="readonly")
        self.table_font.grid(row=1, column=1, sticky="w")
        self.table_font.set("Calibri")
        ttk.Label(fmt_frame, text="Tamaño:").grid(row=1, column=2, sticky="w")
        self.table_size = ttk.Spinbox(fmt_frame, from_=8, to=24, width=5)
        self.table_size.set("11")
        self.table_size.grid(row=1, column=3, sticky="w")

        row += 1
        ttk.Separator(inner, orient="horizontal").grid(row=row, column=0, columnspan=3, sticky="ew", pady=8)
        row += 1

        # Mapeos (Excel -> Word)
        ttk.Label(inner, text="Mapeos Excel ↔ Word", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky="w")
        row += 1

        mapping_frame = ttk.Frame(inner)
        mapping_frame.grid(row=row, column=0, columnspan=3, sticky="ew")
        mapping_frame.columnconfigure(0, weight=1)

        self.mapping_tree = ttk.Treeview(mapping_frame, columns=("encabezado", "hoja", "celda", "tipo"), show="headings", height=6)
        self.mapping_tree.heading("encabezado", text="Encabezado Word")
        self.mapping_tree.heading("hoja", text="Hoja Excel")
        self.mapping_tree.heading("celda", text="Celda/Rango")
        self.mapping_tree.heading("tipo", text="Tipo")
        self.mapping_tree.grid(row=0, column=0, sticky="nsew")
        map_scroll = ttk.Scrollbar(mapping_frame, orient="vertical", command=self.mapping_tree.yview)
        map_scroll.grid(row=0, column=1, sticky="ns")
        self.mapping_tree.configure(yscrollcommand=map_scroll.set)

        row += 1
        # Controls for mapping
        controls = ttk.LabelFrame(inner, text="Agregar/Editar Mapeo", padding=6)
        controls.grid(row=row, column=0, columnspan=3, sticky="ew", pady=6)
        controls.columnconfigure(1, weight=1)
        ttk.Label(controls, text="Encabezado Word:").grid(row=0, column=0, sticky="w")
        self.entry_word_col = ttk.Entry(controls, width=30)
        self.entry_word_col.grid(row=0, column=1, sticky="w")
        ttk.Label(controls, text="Hoja Excel:").grid(row=0, column=2, sticky="w")
        self.entry_sheet = ttk.Entry(controls, width=18)
        self.entry_sheet.grid(row=0, column=3, sticky="w")

        ttk.Label(controls, text="Celda/Rango:").grid(row=1, column=0, sticky="w")
        self.entry_cell = ttk.Entry(controls, width=30)
        self.entry_cell.grid(row=1, column=1, sticky="w")
        ttk.Label(controls, text="Tipo:").grid(row=1, column=2, sticky="w")
        self.combo_type = ttk.Combobox(controls, values=["valor", "promedio"], width=15, state="readonly")
        self.combo_type.grid(row=1, column=3, sticky="w")
        self.combo_type.set("valor")

        ttk.Label(controls, text="Formatos válidos: A1 | A1:A10 | C5,E7,D9").grid(row=2, column=0, columnspan=4, sticky="w", pady=4)

        row += 1
        btns = ttk.Frame(inner)
        btns.grid(row=row, column=0, columnspan=3, pady=6)
        ttk.Button(btns, text="➕ Agregar Mapeo", command=self.add_mapping).pack(side="left", padx=4)
        ttk.Button(btns, text="✏️ Editar Seleccionado", command=self.edit_selected_mapping).pack(side="left", padx=4)
        ttk.Button(btns, text="❌ Eliminar Seleccionado", command=self.delete_selected_mapping).pack(side="left", padx=4)

        row += 1
        ttk.Separator(inner, orient="horizontal").grid(row=row, column=0, columnspan=3, sticky="ew", pady=8)
        row += 1

        # Reemplazos de texto (ordenados)
        ttk.Label(inner, text="Reemplazos de Texto (se ejecutan en el orden listado)", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky="w")
        row += 1

        replace_frame = ttk.Frame(inner)
        replace_frame.grid(row=row, column=0, columnspan=3, sticky="ew")
        self.replace_tree = ttk.Treeview(replace_frame, columns=("original", "nuevo"), show="headings", height=5)
        self.replace_tree.heading("original", text="Texto original")
        self.replace_tree.heading("nuevo", text="Texto nuevo")
        self.replace_tree.grid(row=0, column=0, sticky="nsew")
        rep_scroll = ttk.Scrollbar(replace_frame, orient="vertical", command=self.replace_tree.yview)
        rep_scroll.grid(row=0, column=1, sticky="ns")
        self.replace_tree.configure(yscrollcommand=rep_scroll.set)

        row += 1
        rc = ttk.LabelFrame(inner, text="Agregar Reemplazo", padding=6)
        rc.grid(row=row, column=0, columnspan=3, sticky="ew", pady=6)
        ttk.Label(rc, text="Original:").grid(row=0, column=0, sticky="w")
        self.entry_texto_original = ttk.Entry(rc, width=30)
        self.entry_texto_original.grid(row=0, column=1, sticky="w")
        ttk.Label(rc, text="Nuevo:").grid(row=0, column=2, sticky="w")
        self.entry_texto_nuevo = ttk.Entry(rc, width=30)
        self.entry_texto_nuevo.grid(row=0, column=3, sticky="w")
        ttk.Button(rc, text="➕ Agregar", command=self.add_replacement).grid(row=1, column=0, pady=6)
        ttk.Button(rc, text="❌ Eliminar Seleccionado", command=self.delete_selected_replacement).grid(row=1, column=1, pady=6)

        row += 1
        ttk.Separator(inner, orient="horizontal").grid(row=row, column=0, columnspan=3, sticky="ew", pady=8)
        row += 1

        # Imagenes config
        ttk.Label(inner, text="Imágenes (map. automático por subcarpeta)", font=("Arial", 10, "bold")).grid(row=row, column=0, sticky="w")
        row += 1
        ttk.Label(inner, text="Carpeta de imágenes:").grid(row=row, column=0, sticky="w")
        self.imagenes_folder_entry = ttk.Entry(inner, width=70)
        self.imagenes_folder_entry.grid(row=row, column=1, padx=6, sticky="w")
        ttk.Button(inner, text="📂", width=3, command=self.select_imagenes_folder).grid(row=row, column=2)
        row += 1
        ttk.Label(inner, text="Altura fija (cm):").grid(row=row, column=0, sticky="w")
        self.entry_fixed_height = ttk.Entry(inner, width=10)
        self.entry_fixed_height.grid(row=row, column=1, sticky="w")
        self.entry_fixed_height.insert(0, "5.0")
        row += 1
        self.usar_mapeo_imagenes_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(inner, text="Usar mapeo automático (subcarpetas numeradas)", variable=self.usar_mapeo_imagenes_var, command=self.toggle_mapeo_imagenes).grid(row=row, column=0, columnspan=3, sticky="w")
        row += 1

        self.imagen_mapeo_frame = ttk.LabelFrame(inner, text="Mapeo de Imágenes", padding=6)
        self.imagen_mapeo_frame.grid(row=row, column=0, columnspan=3, sticky="ew")
        self.imagen_mapeo_frame.columnconfigure(0, weight=1)

        self.imagen_tree = ttk.Treeview(self.imagen_mapeo_frame, columns=("posicion", "accion"), show="headings", height=5)
        self.imagen_tree.heading("posicion", text="Posición en Word")
        self.imagen_tree.heading("accion", text="Acción / Subcarpeta")
        self.imagen_tree.grid(row=0, column=0, sticky="nsew")
        im_scroll = ttk.Scrollbar(self.imagen_mapeo_frame, orient="vertical", command=self.imagen_tree.yview)
        im_scroll.grid(row=0, column=1, sticky="ns")
        self.imagen_tree.configure(yscrollcommand=im_scroll.set)

        # controls for image mapping
        ic = ttk.Frame(self.imagen_mapeo_frame)
        ic.grid(row=1, column=0, columnspan=2, pady=6, sticky="w")
        ttk.Label(ic, text="Posición en Word (1..):").grid(row=0, column=0, sticky="w")
        self.imagen_posicion_spin = ttk.Spinbox(ic, from_=1, to=100, width=6)
        self.imagen_posicion_spin.grid(row=0, column=1, sticky="w", padx=4)
        ttk.Label(ic, text="Subcarpeta (número):").grid(row=0, column=2, sticky="w")
        self.imagen_subcarpeta_spin = ttk.Spinbox(ic, from_=1, to=200, width=6)
        self.imagen_subcarpeta_spin.grid(row=0, column=3, sticky="w", padx=4)
        ttk.Button(ic, text="➕ Agregar Mapeo Imagen", command=self.add_imagen_mapping).grid(row=0, column=4, padx=6)
        ttk.Button(ic, text="❌ Eliminar Seleccionado", command=self.delete_selected_imagen_mapping).grid(row=0, column=5, padx=6)
        ttk.Button(ic, text="🔍 Analizar Word", command=self.analyze_word_images).grid(row=0, column=6, padx=6)

        # ocultar por defecto
        if not self.usar_mapeo_imagenes_var.get():
            self.imagen_mapeo_frame.grid_remove()

        # === Right panel (process + logs) ===
        right = ttk.LabelFrame(main, text="Procesamiento y registros", padding=8)
        right.grid(row=1, column=1, sticky="nsew")
        right.columnconfigure(0, weight=1)
        right.rowconfigure(2, weight=1)

        # Config buttons
        cfg = ttk.Frame(right)
        cfg.grid(row=0, column=0, sticky="ew")
        ttk.Button(cfg, text="💾 Guardar Configuración", command=self.guardar_config_json).pack(side="left", padx=3)
        ttk.Button(cfg, text="📁 Cargar Configuración", command=self.cargar_config_json).pack(side="left", padx=3)
        ttk.Button(cfg, text="🗑️ Limpiar Todo", command=self.clear_all_config).pack(side="left", padx=3)

        # Processing controls
        proc = ttk.Frame(right)
        proc.grid(row=1, column=0, sticky="ew", pady=6)
        proc.columnconfigure(0, weight=1)
        ttk.Label(proc, text="Rango de calicatas:").grid(row=0, column=0, sticky="w")
        ttk.Label(proc, text="Desde C-").grid(row=0, column=1, sticky="w")
        self.start_range = ttk.Spinbox(proc, from_=1, to=999, width=6)
        self.start_range.set("1")
        self.start_range.grid(row=0, column=2, sticky="w")
        ttk.Label(proc, text="Hasta C-").grid(row=0, column=3, sticky="w", padx=(6,0))
        self.end_range = ttk.Spinbox(proc, from_=1, to=999, width=6)
        self.end_range.set("10")
        self.end_range.grid(row=0, column=4, sticky="w")

        self.progress = ttk.Progressbar(proc, orient="horizontal", mode="determinate")
        self.progress.grid(row=1, column=0, columnspan=5, sticky="ew", pady=6)
        self.progress_info = ttk.Label(proc, text="")
        self.progress_info.grid(row=2, column=0, columnspan=5, sticky="w")

        pb = ttk.Frame(proc)
        pb.grid(row=3, column=0, columnspan=5, pady=4)
        self.process_btn = ttk.Button(pb, text="🚀 Procesar Informes", command=self.run_processing_threaded)
        self.process_btn.pack(side="left", padx=4)
        self.stop_btn = ttk.Button(pb, text="⏹️ Detener", command=self.stop_processing, state="disabled")
        self.stop_btn.pack(side="left", padx=4)
        ttk.Button(pb, text="📂 Abrir Carpeta Salida", command=self.open_output_folder).pack(side="left", padx=4)

        # Logs
        ttk.Label(right, text="Registro de actividad").grid(row=2, column=0, sticky="w")
        logf = ttk.Frame(right)
        logf.grid(row=3, column=0, sticky="nsew")
        logf.columnconfigure(0, weight=1)
        logf.rowconfigure(0, weight=1)
        self.log_console = tk.Text(logf, height=15, wrap="word")
        self.log_console.grid(row=0, column=0, sticky="nsew")
        log_scroll = ttk.Scrollbar(logf, orient="vertical", command=self.log_console.yview)
        log_scroll.grid(row=0, column=1, sticky="ns")
        self.log_console.configure(yscrollcommand=log_scroll.set)

        lbtns = ttk.Frame(right)
        lbtns.grid(row=4, column=0, sticky="w", pady=6)
        ttk.Button(lbtns, text="🗑️ Limpiar Log", command=self.clear_log).pack(side="left", padx=4)
        ttk.Button(lbtns, text="💾 Guardar Log", command=self.save_log).pack(side="left", padx=4)

        # Inicializaciones rápidas
        self.update_status("Listo", "#27ae60")

    # -------------------------
    # Métodos UI / utilitarios
    # -------------------------

    def log(self, text):
        ts = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        msg = f"[{ts}] {text}\n"
        try:
            self.log_console.insert("end", msg)
            self.log_console.see("end")
        except Exception:
            print(msg, end="")

    def clear_log(self):
        self.log_console.delete("1.0", "end")

    def save_log(self):
        content = self.log_console.get("1.0", "end")
        if not content.strip():
            messagebox.showinfo("Guardar log", "No hay contenido en el log.")
            return
        path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text files","*.txt")])
        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    f.write(content)
                messagebox.showinfo("Guardado", f"Log guardado en:\n{path}")
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo guardar el log:\n{str(e)}")

    def update_status(self, text, color=None):
        try:
            self.status_label.config(text=text)
            if color:
                self.status_label.config(foreground=color)
        except Exception:
            pass

    # -------------------------
    # Validaciones y helpers
    # -------------------------

    def validate_cell_format_enhanced(self, celltext):
        """
        Validar formatos aceptados: A1, A1:A10, C5:E7, C5,D7,F9, etc.
        Regla básica: columnas (A..ZZ) y filas numéricas.
        """
        if not celltext or not isinstance(celltext, str):
            return False
        parts = [p.strip() for p in celltext.split(",")]
        pattern_single = re.compile(r"^[A-Z]{1,3}\d{1,6}$", re.I)
        pattern_range = re.compile(r"^[A-Z]{1,3}\d{1,6}:[A-Z]{1,3}\d{1,6}$", re.I)
        for p in parts:
            if pattern_single.match(p):
                continue
            if pattern_range.match(p):
                continue
            return False
        return True

    # -------------------------
    # Mapeos (agregar/editar/eliminar)
    # -------------------------

    def add_mapping(self):
        word_col = self.entry_word_col.get().strip()
        sheet = self.entry_sheet.get().strip()
        cell = self.entry_cell.get().strip()
        tipo = self.combo_type.get().strip()
        if not all([word_col, sheet, cell, tipo]):
            messagebox.showerror("Error", "Complete todos los campos del mapeo.")
            return
        if not self.validate_cell_format_enhanced(cell):
            messagebox.showerror("Error", "Formato de celda inválido. Ej: A1 | A1:A10 | C5,E7")
            return

        if self.editing_item:
            # actualizar
            self.mapping_tree.item(self.editing_item, values=(word_col, sheet, cell, tipo))
            idx = list(self.mapping_tree.get_children()).index(self.editing_item)
            self.config["mappings"][idx] = {"encabezado": word_col, "hoja": sheet, "celda": cell, "tipo": tipo}
            self.log(f"✏️ Mapeo editado: {word_col} ← {sheet}[{cell}] ({tipo})")
            self.cancel_edit_mapping()
        else:
            self.mapping_tree.insert("", "end", values=(word_col, sheet, cell, tipo))
            self.config["mappings"].append({"encabezado": word_col, "hoja": sheet, "celda": cell, "tipo": tipo})
            self.log(f"➕ Mapeo agregado: {word_col} ← {sheet}[{cell}] ({tipo})")

        self.clear_mapping_fields()

    def clear_mapping_fields(self):
        self.entry_word_col.delete(0, "end")
        self.entry_sheet.delete(0, "end")
        self.entry_cell.delete(0, "end")
        self.combo_type.set("valor")
        self.editing_item = None

    def edit_selected_mapping(self):
        sel = self.mapping_tree.selection()
        if not sel:
            messagebox.showwarning("Advertencia", "Seleccione un mapeo para editar.")
            return
        item = sel[0]
        vals = self.mapping_tree.item(item)["values"]
        if not vals:
            return
        self.entry_word_col.delete(0, "end"); self.entry_word_col.insert(0, vals[0])
        self.entry_sheet.delete(0, "end"); self.entry_sheet.insert(0, vals[1])
        self.entry_cell.delete(0, "end"); self.entry_cell.insert(0, vals[2])
        self.combo_type.set(vals[3])
        self.editing_item = item

    def delete_selected_mapping(self):
        sel = self.mapping_tree.selection()
        if not sel:
            messagebox.showwarning("Advertencia", "Seleccione mapeos a eliminar.")
            return
        if not messagebox.askyesno("Confirmar", "¿Eliminar mapeos seleccionados?"):
            return
        for it in sel:
            idx = list(self.mapping_tree.get_children()).index(it)
            try:
                del self.config["mappings"][idx]
            except Exception:
                pass
            self.mapping_tree.delete(it)
        self.log("❌ Mapeo(s) eliminado(s).")

    # -------------------------
    # Reemplazos de texto
    # -------------------------

    def add_replacement(self):
        orig = self.entry_texto_original.get().strip()
        nuevo = self.entry_texto_nuevo.get().strip()
        if orig == "":
            messagebox.showerror("Error", "Texto original vacío.")
            return
        self.replace_tree.insert("", "end", values=(orig, nuevo))
        self.config["text_replacements"].append((orig, nuevo))
        self.entry_texto_original.delete(0, "end")
        self.entry_texto_nuevo.delete(0, "end")
        self.log(f"➕ Reemplazo agregado: '{orig}' → '{nuevo}'")

    def delete_selected_replacement(self):
        sel = self.replace_tree.selection()
        if not sel:
            messagebox.showwarning("Advertencia", "Seleccione un reemplazo para eliminar.")
            return
        if not messagebox.askyesno("Confirmar", "¿Eliminar reemplazos seleccionados?"):
            return
        for it in sel:
            idx = list(self.replace_tree.get_children()).index(it)
            try:
                del self.config["text_replacements"][idx]
            except Exception:
                pass
            self.replace_tree.delete(it)
        self.log("❌ Reemplazo(s) eliminado(s).")

    # -------------------------
    # Imágenes (config y procesamiento)
    # -------------------------

    def toggle_mapeo_imagenes(self):
        if self.usar_mapeo_imagenes_var.get():
            self.imagen_mapeo_frame.grid()
        else:
            self.imagen_mapeo_frame.grid_remove()

    def add_imagen_mapping(self):
        pos = int(self.imagen_posicion_spin.get())
        sub = int(self.imagen_subcarpeta_spin.get())
        desc = f"Subcarpeta #{sub}"
        self.imagen_tree.insert("", "end", values=(f"Imagen {pos}", desc))
        self.config["imagen_config"]["imagen_mapeos"].append({"posicion": pos, "imagen_subcarpeta": sub})
        self.log(f"➕ Mapeo imagen agregado: pos {pos} ← subcarpeta {sub}")

    def delete_selected_imagen_mapping(self):
        sel = self.imagen_tree.selection()
        if not sel:
            messagebox.showwarning("Advertencia", "Seleccione mapeo(s) de imagen para eliminar.")
            return
        if not messagebox.askyesno("Confirmar", "¿Eliminar mapeos de imagen seleccionados?"):
            return
        for it in sel:
            idx = list(self.imagen_tree.get_children()).index(it)
            try:
                del self.config["imagen_config"]["imagen_mapeos"][idx]
            except Exception:
                pass
            self.imagen_tree.delete(it)
        self.log("❌ Mapeo(s) de imagen eliminado(s).")

    def analyze_word_images(self):
        """Analiza el documento base y muestra la cantidad de imágenes y posiciones detectadas."""
        path = self.docx_entry.get().strip()
        if not path or not os.path.exists(path):
            messagebox.showwarning("Advertencia", "Seleccione un documento Word válido para analizar.")
            return
        try:
            doc = Document(path)
            imgs = listar_imagenes_doc(doc)
            # limpiar tree
            for it in self.imagen_tree.get_children():
                self.imagen_tree.delete(it)
            for i, img in enumerate(imgs, start=1):
                self.imagen_tree.insert("", "end", values=(f"Imagen {i}", "No configurado"))
            self.log(f"🔍 Análisis completado: {len(imgs)} imágenes encontradas.")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo analizar el documento: {str(e)}")
            self.log(f"❌ Error analizando documento: {str(e)}")

    def seleccionar_imagen_por_subcarpeta(self, root_folder, sub_num):
        """
        Busca subcarpeta que contenga el número indicado (ej: '01', 'C-01', 'Imágenes 01').
        Si no encuentra, intenta seleccionar por orden de modificación.
        """
        if not os.path.exists(root_folder):
            return None
        # buscar subcarpeta con número
        entries = [d for d in os.listdir(root_folder) if os.path.isdir(os.path.join(root_folder, d))]
        # prefer exact numeric suffixes or contains numeric sequence
        pattern = re.compile(r"(\d{1,3})")
        for d in entries:
            m = pattern.search(d)
            if m:
                if int(m.group(1)) == sub_num:
                    return os.path.join(root_folder, d)
        # fallback: orden por fecha
        try:
            entries_full = [os.path.join(root_folder, d) for d in entries]
            entries_full.sort(key=lambda x: os.path.getmtime(x))
            idx = sub_num - 1
            if idx < len(entries_full):
                return entries_full[idx]
        except Exception:
            pass
        return None

    def obtener_imagenes_ordenadas(self, subcarpeta):
        exts = ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff']
        if not subcarpeta or not os.path.exists(subcarpeta):
            return []
        files = [os.path.join(subcarpeta, f) for f in os.listdir(subcarpeta) if any(f.lower().endswith(e) for e in exts)]
        files.sort(key=lambda x: os.path.getmtime(x))
        return files

    def procesar_imagenes_calicata(self, doc, calicata, numero):
        """
        Reemplazos según mapeo automático (imagen_config.imagen_mapeos):
          - Busca la subcarpeta correspondiente (por número).
          - Reemplaza las imágenes por orden (más antigua -> primero).
        """
        root = self.imagenes_folder_entry.get().strip() or self.config.get("imagenes_folder", "")
        if not root or not os.path.exists(root):
            self.log(f"⚠️ No hay carpeta de imágenes configurada.")
            return
        subcarpeta = self.seleccionar_imagen_por_subcarpeta(root, numero)
        if not subcarpeta or not os.path.exists(subcarpeta):
            self.log(f"⚠️ No se encontró subcarpeta para {calicata} (buscando {numero}).")
            return
        imgs = self.obtener_imagenes_ordenadas(subcarpeta)
        if not imgs:
            self.log(f"⚠️ Subcarpeta {os.path.basename(subcarpeta)} no contiene imágenes.")
            return

        imgs_doc = listar_imagenes_doc(doc)
        for m in self.config["imagen_config"].get("imagen_mapeos", []):
            pos = m.get("posicion", 1) - 1
            subidx = m.get("imagen_subcarpeta", 1) - 1
            # elegir imagen en la subcarpeta según índice
            if subidx < 0 or subidx >= len(imgs):
                self.log(f"⚠️ Sub índice {subidx+1} fuera de rango en subcarpeta {subcarpeta}")
                continue
            if pos < 0 or pos >= len(imgs_doc):
                self.log(f"⚠️ Posición imagen {pos+1} no encontrada en el documento")
                continue
            ruta_nueva = imgs[subidx]
            info = imgs_doc[pos]
            try:
                fixed_h = None
                try:
                    fixed_h = float(self.entry_fixed_height.get() or self.config.get("fixed_image_height", 5.0))
                except Exception:
                    fixed_h = self.config.get("fixed_image_height", 5.0)
                reemplazar_imagen(info["run"], ruta_nueva, fixed_h)
                self.log(f"🖼️ Imagen {pos+1} reemplazada por {os.path.basename(ruta_nueva)}")
            except Exception as e:
                self.log(f"⚠️ Error reemplazando imagen {pos+1}: {str(e)}")

    def procesar_imagenes_consolidado(self, doc, start_val, end_val):
        # toma la primera calicata como ejemplo
        return self.procesar_imagenes_calicata(doc, f"C-{start_val:02d}", start_val)

    # -------------------------
    # Búsqueda archivo Excel
    # -------------------------

    def buscar_archivo_excel(self, calicata):
        """Buscar calicata.xlsx en carpetas configuradas"""
        for carpeta in [self.excel_folder_entry_1.get().strip(), self.excel_folder_entry_2.get().strip(), self.config.get("excel_folder_1",""), self.config.get("excel_folder_2","")]:
            if carpeta and os.path.exists(carpeta):
                candidate = os.path.join(carpeta, f"{calicata}.xlsx")
                if os.path.exists(candidate):
                    return candidate
        return None

    # -------------------------
    # Guardar/Cargar/Reset config
    # -------------------------

    def save_config(self):
        """Guardar en self.config los valores actuales del UI"""
        self.config["docx_path"] = self.docx_entry.get().strip()
        self.config["excel_folder_1"] = self.excel_folder_entry_1.get().strip()
        self.config["excel_folder_2"] = self.excel_folder_entry_2.get().strip()
        self.config["output_folder"] = self.output_folder_entry.get().strip()
        self.config["imagenes_folder"] = self.imagenes_folder_entry.get().strip()
        try:
            self.config["fixed_image_height"] = float(self.entry_fixed_height.get())
        except Exception:
            self.config["fixed_image_height"] = 5.0

        self.config["font_config"] = {
            "paragraph_font": self.paragraph_font.get(),
            "paragraph_size": int(self.paragraph_size.get()),
            "table_font": self.table_font.get(),
            "table_size": int(self.table_size.get())
        }
        self.config["archivo_config"] = {
            "nombre_base": self.nombre_base_entry.get().strip(),
            "usar_sufijo": bool(self.usar_sufijo_var.get()),
            "sufijo_personalizado": self.sufijo_entry.get().strip()
        }
        self.config["imagen_config"]["usar_mapeo_automatico"] = bool(self.usar_mapeo_imagenes_var.get())
        self.config["informe_config"] = {
            "tipo_informe": self.tipo_informe_var.get(),
            "consolidado_nombre": self.consolidado_nombre_entry.get().strip()
        }
        self.config["decimales_config"] = {
            "usar_decimales_fijos": bool(self.usar_decimales_var.get()),
            "cantidad_decimales": int(self.decimales_spin.get())
        }
        # mappings, replacements and imagen_mapeos are updated as user modifies trees (we kept them in sync earlier)

    def guardar_config_json(self):
        self.save_config()
        fname = f"config_calicatas_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        path = filedialog.asksaveasfilename(defaultextension=".json", initialfile=fname, filetypes=[("JSON","*.json")])
        if not path:
            return
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(self.config, f, indent=2, ensure_ascii=False)
            self.last_config_file = path
            self.log(f"💾 Config guardada: {os.path.basename(path)}")
            messagebox.showinfo("Guardado", f"Configuración guardada en:\n{path}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar la configuración:\n{str(e)}")
            self.log(f"❌ Error guardando config: {str(e)}")

    def cargar_config_json(self):
        path = filedialog.askopenfilename(filetypes=[("JSON","*.json")])
        if not path:
            return
        try:
            with open(path, "r", encoding="utf-8") as f:
                cfg = json.load(f)
            self.config.update(cfg)
            # cargar UI fields (asegurar claves)
            self.docx_entry.delete(0, "end"); self.docx_entry.insert(0, self.config.get("docx_path",""))
            self.excel_folder_entry_1.delete(0,"end"); self.excel_folder_entry_1.insert(0, self.config.get("excel_folder_1",""))
            self.excel_folder_entry_2.delete(0,"end"); self.excel_folder_entry_2.insert(0, self.config.get("excel_folder_2",""))
            self.output_folder_entry.delete(0,"end"); self.output_folder_entry.insert(0, self.config.get("output_folder",""))
            self.imagenes_folder_entry.delete(0,"end"); self.imagenes_folder_entry.insert(0, self.config.get("imagenes_folder",""))
            self.entry_fixed_height.delete(0,"end"); self.entry_fixed_height.insert(0, str(self.config.get("fixed_image_height",5.0)))
            # fuentes
            fc = self.config.get("font_config",{})
            self.paragraph_font.set(fc.get("paragraph_font","Calibri"))
            self.paragraph_size.set(str(fc.get("paragraph_size",11)))
            self.table_font.set(fc.get("table_font","Calibri"))
            self.table_size.set(str(fc.get("table_size",11)))
            # archivo cfg
            ac = self.config.get("archivo_config",{})
            self.nombre_base_entry.delete(0,"end"); self.nombre_base_entry.insert(0, ac.get("nombre_base","EMS CUSCO C-"))
            self.usar_sufijo_var.set(ac.get("usar_sufijo", True))
            self.sufijo_entry.config(state="normal"); self.sufijo_entry.delete(0,"end"); self.sufijo_entry.insert(0, ac.get("sufijo_personalizado",""))
            if ac.get("usar_sufijo", True):
                self.sufijo_entry.config(state="disabled")
            # informe
            inf = self.config.get("informe_config",{})
            self.tipo_informe_var.set(inf.get("tipo_informe","individual"))
            self.consolidado_nombre_entry.delete(0,"end"); self.consolidado_nombre_entry.insert(0, inf.get("consolidado_nombre","Informe_Consolidado"))
            # decimales
            dec = self.config.get("decimales_config",{})
            self.usar_decimales_var.set(dec.get("usar_decimales_fijos", True))
            self.decimales_spin.set(str(dec.get("cantidad_decimales",1)))
            # limpiar trees
            for t in (self.mapping_tree, self.replace_tree, self.imagen_tree):
                for it in t.get_children():
                    t.delete(it)
            # cargar mappings
            for m in self.config.get("mappings",[]):
                self.mapping_tree.insert("", "end", values=(m.get("encabezado",""), m.get("hoja",""), m.get("celda",""), m.get("tipo","valor")))
            for r in self.config.get("text_replacements", []):
                self.replace_tree.insert("", "end", values=(r[0], r[1]))
            for im in self.config.get("imagen_config", {}).get("imagen_mapeos", []):
                self.imagen_tree.insert("", "end", values=(f"Imagen {im.get('posicion')}", f"Subcarpeta {im.get('imagen_subcarpeta')}"))
            self.last_config_file = path
            self.log(f"📁 Config cargada: {os.path.basename(path)}")
            messagebox.showinfo("Cargado", f"Configuración cargada desde:\n{path}")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo cargar la configuración:\n{str(e)}")
            self.log(f"❌ Error cargando config: {str(e)}")

    def clear_all_config(self):
        if not messagebox.askyesno("Confirmar", "¿Limpiar toda la configuración y mapeos?"):
            return
        # limpiar UI campos
        entries = [self.docx_entry, self.excel_folder_entry_1, self.excel_folder_entry_2, self.output_folder_entry, self.imagenes_folder_entry]
        for e in entries:
            e.delete(0, "end")
        self.nombre_base_entry.delete(0,"end"); self.nombre_base_entry.insert(0, "EMS CUSCO C-")
        self.sufijo_entry.config(state="disabled"); self.sufijo_entry.delete(0,"end")
        self.consolidado_nombre_entry.delete(0,"end"); self.consolidado_nombre_entry.insert(0, "Informe_Consolidado")
        self.entry_fixed_height.delete(0,"end"); self.entry_fixed_height.insert(0, "5.0")
        # limpiar árboles
        for t in (self.mapping_tree, self.replace_tree, self.imagen_tree):
            for it in t.get_children():
                t.delete(it)
        # reset config
        self.config = {
            "docx_path": "",
            "excel_folder_1": "",
            "excel_folder_2": "",
            "output_folder": "",
            "mappings": [],
            "text_replacements": [],
            "imagenes_folder": "",
            "image_replacements": [],
            "fixed_image_height": 5.0,
            "font_config": {"paragraph_font":"Calibri", "paragraph_size":11, "table_font":"Calibri", "table_size":11},
            "archivo_config": {"nombre_base":"EMS CUSCO C-", "usar_sufijo":True, "sufijo_personalizado":""},
            "imagen_config": {"usar_mapeo_automatico": False, "imagen_mapeos": []},
            "informe_config": {"tipo_informe": "individual", "consolidado_nombre": "Informe_Consolidado"},
            "decimales_config": {"usar_decimales_fijos": True, "cantidad_decimales": 1}
        }
        self.log("🗑️ Configuración limpiada.")

    # -------------------------
    # Selección de rutas (UI)
    # -------------------------

    def select_docx(self):
        path = filedialog.askopenfilename(filetypes=[("Word Files","*.docx")])
        if path:
            self.docx_entry.delete(0,"end"); self.docx_entry.insert(0, path)
            self.log(f"📄 Documento seleccionado: {os.path.basename(path)}")

    def select_excel_folder(self, num=1):
        path = filedialog.askdirectory()
        if path:
            if num == 1:
                self.excel_folder_entry_1.delete(0,"end"); self.excel_folder_entry_1.insert(0, path)
            else:
                self.excel_folder_entry_2.delete(0,"end"); self.excel_folder_entry_2.insert(0, path)
            # contar archivos
            try:
                files = [f for f in os.listdir(path) if f.lower().endswith(".xlsx")]
                self.log(f"📁 Carpeta Excel {num} seleccionada: {len(files)} archivos .xlsx")
            except Exception:
                pass

    def select_output_folder(self):
        path = filedialog.askdirectory()
        if path:
            self.output_folder_entry.delete(0,"end"); self.output_folder_entry.insert(0, path)
            self.log(f"📂 Carpeta de salida: {path}")

    def select_imagenes_folder(self):
        path = filedialog.askdirectory()
        if path:
            self.imagenes_folder_entry.delete(0,"end"); self.imagenes_folder_entry.insert(0, path)
            self.log(f"🖼️ Carpeta de imágenes: {path}")

    def toggle_sufijo(self):
        if self.usar_sufijo_var.get():
            self.sufijo_entry.config(state="disabled")
        else:
            self.sufijo_entry.config(state="normal")

    def open_output_folder(self):
        path = self.output_folder_entry.get().strip()
        if path and os.path.exists(path):
            open_path(path)
        else:
            messagebox.showwarning("Advertencia", "Carpeta de salida inválida.")

    # -------------------------
    # Procesamiento principal
    # -------------------------

    def run_processing_threaded(self):
        # validar configuración mínima
        errors = self.validate_config()
        if errors:
            messagebox.showerror("Error de configuración", "Errores:\n" + "\n".join(errors))
            return
        if self.processing:
            messagebox.showwarning("Advertencia", "Procesamiento ya en curso.")
            return

        start_val = int(self.start_range.get())
        end_val = int(self.end_range.get())
        if start_val > end_val:
            messagebox.showerror("Error", "Rango inválido.")
            return
        count = end_val - start_val + 1
        tipo = self.tipo_informe_var.get()
        if tipo == "individual":
            msg = f"¿Procesar {count} calicatas individuales (C-{start_val:02d} a C-{end_val:02d})?"
        else:
            msg = f"¿Crear informe consolidado con {count} calicatas (C-{start_val:02d} a C-{end_val:02d})?"
        if not messagebox.askyesno("Confirmar", msg):
            return

        self.stop_processing_flag = False
        self.processing = True
        self.process_btn.config(state="disabled")
        self.stop_btn.config(state="normal")
        # guardar config actual
        self.save_config()
        # start thread
        t = threading.Thread(target=self.run_processing, daemon=True)
        t.start()

    def validate_config(self):
        errs = []
        if not self.docx_entry.get().strip() or not os.path.exists(self.docx_entry.get().strip()):
            errs.append("- Seleccione un documento Word base válido.")
        if not (self.excel_folder_entry_1.get().strip() or self.excel_folder_entry_2.get().strip()):
            errs.append("- Seleccione al menos una carpeta de Excel.")
        if not self.output_folder_entry.get().strip() or not os.path.exists(self.output_folder_entry.get().strip()):
            errs.append("- Seleccione una carpeta de salida válida.")
        if not self.mapping_tree.get_children():
            errs.append("- Configure al menos un mapeo Excel ↔ Word.")
        return errs

    def stop_processing(self):
        self.stop_processing_flag = True
        self.update_status("Deteniendo...", "#f39c12")
        self.log("⏹️ Solicitado detener procesamiento...")

    def run_processing(self):
        try:
            self.update_status("Procesando...", "#f39c12")
            tipo = self.tipo_informe_var.get()
            start_val = int(self.start_range.get())
            end_val = int(self.end_range.get())
            if tipo == "individual":
                self.procesar_informes_individuales(start_val, end_val)
            else:
                self.procesar_informe_consolidado(start_val, end_val)
        except Exception as e:
            self.log(f"❌ Error crítico: {str(e)}")
            messagebox.showerror("Error crítico", str(e))
            self.update_status("Error", "#e74c3c")
        finally:
            self.processing = False
            self.process_btn.config(state="normal")
            self.stop_btn.config(state="disabled")
            self.progress["value"] = 0

    def generar_nombre_archivo(self, numero):
        base = self.nombre_base_entry.get().strip() or self.config["archivo_config"]["nombre_base"]
        usar_suf = self.usar_sufijo_var.get()
        if usar_suf:
            return f"{base}{numero:02d}"
        else:
            s = self.sufijo_entry.get().strip()
            return f"{base}{s}" if s else base

    def procesar_informes_individuales(self, start_val, end_val):
        total = end_val - start_val + 1
        self.progress["maximum"] = total
        processed = 0
        errors = 0
        self.log(f"🚀 Iniciando procesamiento individual: {start_val}..{end_val}")
        # ensure config mapping list in sync
        # update config mappings from tree in case user changed and didn't save
        self.config["mappings"] = []
        for it in self.mapping_tree.get_children():
            v = self.mapping_tree.item(it)["values"]
            self.config["mappings"].append({"encabezado": v[0], "hoja": v[1], "celda": v[2], "tipo": v[3]})

        # update replacements
        self.config["text_replacements"] = []
        for it in self.replace_tree.get_children():
            v = self.replace_tree.item(it)["values"]
            self.config["text_replacements"].append((v[0], v[1]))

        for i in range(start_val, end_val + 1):
            if self.stop_processing_flag:
                self.log("⏹️ Procesamiento detenido por el usuario.")
                break
            calicata = f"C-{i:02d}"
            self.progress_info.config(text=f"Procesando {calicata} ({i-start_val+1}/{total})")
            try:
                doc = Document(self.docx_entry.get().strip())
                # reemplazar marcador base
                reemplazar_texto_global(doc, "C-01", calicata)

                # aplicar reemplazos en orden
                for old, new in self.config.get("text_replacements", []):
                    reemplazar_texto_global(doc, old, new)

                aplicar_formato_documento(doc, self.config["font_config"])

                excel_path = self.buscar_archivo_excel(calicata)
                if not excel_path:
                    raise FileNotFoundError(f"No se encontró archivo Excel para {calicata}")

                # aplicar mapeos
                for mapping in self.config.get("mappings", []):
                    try:
                        val = extraer_dato_excel_mejorado(excel_path, mapping["hoja"], mapping["celda"], mapping["tipo"], self.config["decimales_config"])
                        insertar_datos_en_tablas_mejorado(doc, mapping["encabezado"], val, self.config["font_config"])
                        self.log(f"  ✅ {mapping['encabezado']} = {val}")
                    except Exception as e:
                        self.log(f"  ⚠️ Error mapeo {mapping.get('encabezado')}: {str(e)}")

                # imágenes
                if self.usar_mapeo_imagenes_var.get():
                    self.config["imagen_config"]["imagen_mapeos"] = self.config.get("imagen_config", {}).get("imagen_mapeos", self.config["imagen_config"].get("imagen_mapeos", []))
                    self.procesar_imagenes_calicata(doc, calicata, i)

                # nombre y guardar
                nombre = self.generar_nombre_archivo(i)
                outpath = os.path.join(self.output_folder_entry.get().strip(), f"{nombre}.docx")
                doc.save(outpath)
                processed += 1
                self.log(f"✅ {calicata} -> {os.path.basename(outpath)}")
            except Exception as e:
                errors += 1
                self.log(f"❌ Error con {calicata}: {str(e)}")
            finally:
                self.progress["value"] = processed + errors
                self.root.update_idletasks()

        # resumen
        if not self.stop_processing_flag:
            if errors == 0:
                self.log(f"🎉 Procesamiento completado: {processed} archivos generados.")
                self.update_status("Completado", "#27ae60")
            else:
                self.log(f"⚠️ Procesamiento finalizó con {errors} errores. {processed} exitosos.")
                self.update_status("Completado con errores", "#f39c12")
        else:
            self.update_status("Detenido", "#e74c3c")

    def procesar_informe_consolidado(self, start_val, end_val):
        total = end_val - start_val + 1
        self.progress["maximum"] = total
        self.progress["value"] = 0
        self.log(f"🚀 Iniciando informe consolidado: {start_val}..{end_val}")

        # preparar mappings y reemplazos desde los trees
        self.config["mappings"] = []
        for it in self.mapping_tree.get_children():
            v = self.mapping_tree.item(it)["values"]
            self.config["mappings"].append({"encabezado": v[0], "hoja": v[1], "celda": v[2], "tipo": v[3]})
        self.config["text_replacements"] = []
        for it in self.replace_tree.get_children():
            v = self.replace_tree.item(it)["values"]
            self.config["text_replacements"].append((v[0], v[1]))

        # abrir doc base
        try:
            doc = Document(self.docx_entry.get().strip())
            aplicar_formato_documento(doc, self.config["font_config"])
            for old, new in self.config.get("text_replacements", []):
                reemplazar_texto_global(doc, old, new)

            datos_consolidados = {}
            for i in range(start_val, end_val + 1):
                if self.stop_processing_flag:
                    self.log("⏹️ Procesamiento detenido por el usuario.")
                    return
                calicata = f"C-{i:02d}"
                self.progress_info.config(text=f"Recopilando {calicata} ({i-start_val+1}/{total})")
                try:
                    excel_path = self.buscar_archivo_excel(calicata)
                    if not excel_path:
                        self.log(f"⚠️ No encontrado Excel para {calicata}")
                        datos_consolidados[calicata] = {}
                        continue
                    datos = {}
                    for mapping in self.config.get("mappings", []):
                        try:
                            val = extraer_dato_excel_mejorado(excel_path, mapping["hoja"], mapping["celda"], mapping["tipo"], self.config["decimales_config"])
                            datos[mapping["encabezado"]] = val
                        except Exception:
                            datos[mapping["encabezado"]] = ""
                    datos_consolidados[calicata] = datos
                except Exception as e:
                    self.log(f"⚠️ Error recopilando {calicata}: {str(e)}")
                    datos_consolidados[calicata] = {}
                self.progress["value"] = i - start_val + 1
                self.root.update_idletasks()

            # insertar en tablas:
            self.insertar_datos_consolidados(doc, datos_consolidados)
            # imágenes
            if self.usar_mapeo_imagenes_var.get():
                self.procesar_imagenes_consolidado(doc, start_val, end_val)
            # guardar
            nombre = self.consolidado_nombre_entry.get().strip() or self.config["informe_config"].get("consolidado_nombre","Informe_Consolidado")
            outpath = os.path.join(self.output_folder_entry.get().strip(), f"{nombre}.docx")
            doc.save(outpath)
            self.log(f"🎉 Informe consolidado guardado: {os.path.basename(outpath)}")
            self.update_status("Completado", "#27ae60")
        except Exception as e:
            self.log(f"❌ Error generando consolidado: {str(e)}")
            self.update_status("Error", "#e74c3c")

    def insertar_datos_consolidados(self, doc, datos_consolidados):
        """
        Inserta valores en las tablas del doc en base a datos_consolidados:
        datos_consolidados = { "C-01": { "EncabezadoFila": valor, ... }, ... }
        El algoritmo intenta encontrar columnas que coincidan con los nombres de calicata.
        """
        try:
            for tabla in doc.tables:
                if not tabla.rows:
                    continue
                headers = [c.text.strip() for c in tabla.rows[0].cells]
                cal_cols = {}
                for idx, head in enumerate(headers):
                    for cal in datos_consolidados.keys():
                        if cal in head or head in cal:
                            cal_cols[idx] = cal
                            break
                for r_idx in range(1, len(tabla.rows)):
                    row = tabla.rows[r_idx]
                    row_header = row.cells[0].text.strip() if row.cells else ""
                    for col_idx, cal in cal_cols.items():
                        if col_idx < len(row.cells):
                            datos = datos_consolidados.get(cal, {})
                            if row_header in datos:
                                val = datos[row_header]
                                row.cells[col_idx].text = str(val)
                                # formato en cell
                                for p in row.cells[col_idx].paragraphs:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in p.runs:
                                        try:
                                            run.font.name = self.config["font_config"]["table_font"]
                                            run.font.size = Pt(self.config["font_config"]["table_size"])
                                        except Exception:
                                            pass
        except Exception as e:
            self.log(f"⚠️ Error insertando datos consolidados: {str(e)}")

# ---------------------------
# Método de compatibilidad para insertar datos en tablas por encabezado
# ---------------------------

def insertar_datos_en_tablas_mejorado(doc, encabezado, valor, font_config=None):
    """
    Busca la columna cuyo encabezado sea exactamente 'encabezado' en la primera fila de cada tabla
    y escribe 'valor' en todas las celdas de esa columna (filas de datos).
    """
    try:
        for tabla in doc.tables:
            if not tabla.rows:
                continue
            headers = [cell.text.strip() for cell in tabla.rows[0].cells]
            for idx, h in enumerate(headers):
                if h == encabezado:
                    for row in tabla.rows[1:]:
                        if idx < len(row.cells):
                            cell = row.cells[idx]
                            cell.text = str(valor)
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in p.runs:
                                    try:
                                        if font_config:
                                            run.font.name = font_config.get("table_font", "Calibri")
                                            run.font.size = Pt(font_config.get("table_size", 11))
                                    except Exception:
                                        pass
                            try:
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            except Exception:
                                pass
                    break
    except Exception:
        pass

# ---------------------------
# MAIN
# ---------------------------

if __name__ == "__main__":
    root = tk.Tk()
    app = CalicataApp(root)
    try:
        root.mainloop()
    except KeyboardInterrupt:
        pass
